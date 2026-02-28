#!/usr/bin/env python3
"""Generate a professional Word document from the PR Review Checklist Template markdown."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "2B579A")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = doc.styles['Code Block']
    p.clear()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    return p


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Callout']
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    return p


def add_checkbox_item(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Checkbox']
    p.clear()
    run = p.add_run('\u2610  ')
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p


def add_checked_item(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Checkbox']
    p.clear()
    run = p.add_run('\u2611  ')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = doc.styles['Custom Bullet']
    p.clear()
    run = p.add_run('\u2022  ')
    run.font.size = Pt(10)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p


def add_numbered_item(doc, number, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = doc.styles['Custom Bullet']
    p.clear()
    run = p.add_run(f'{number}.  ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p


def add_sub_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = doc.styles['Custom Bullet']
    p.paragraph_format.left_indent = Inches(0.75)
    p.clear()
    run = p.add_run('\u2013  ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p


def add_section_label(doc, text):
    """Add a colored section label (e.g., Correctness, Code Quality)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    return p


def create_styles(doc):
    styles = doc.styles

    # Title
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(4)

    # Subtitle
    subtitle = styles['Subtitle']
    subtitle.font.name = 'Calibri'
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    pPr = h1.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="2B579A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Code Block
    code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.right_indent = Inches(0.3)
    pPr = code_style.element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E" w:val="clear"/>')
    pPr.append(shading)

    # Callout
    callout_style = styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH)
    callout_style.font.name = 'Calibri'
    callout_style.font.size = Pt(10)
    callout_style.font.italic = True
    callout_style.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    callout_style.paragraph_format.left_indent = Inches(0.4)
    callout_style.paragraph_format.space_before = Pt(6)
    callout_style.paragraph_format.space_after = Pt(6)
    pPr = callout_style.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="2B579A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF3FA" w:val="clear"/>')
    pPr.append(shading)

    # Checkbox
    checkbox_style = styles.add_style('Checkbox', WD_STYLE_TYPE.PARAGRAPH)
    checkbox_style.font.name = 'Calibri'
    checkbox_style.font.size = Pt(10)
    checkbox_style.paragraph_format.space_before = Pt(2)
    checkbox_style.paragraph_format.space_after = Pt(2)
    checkbox_style.paragraph_format.left_indent = Inches(0.3)

    # Custom Bullet
    bullet = styles.add_style('Custom Bullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet.font.name = 'Calibri'
    bullet.font.size = Pt(10)
    bullet.paragraph_format.space_before = Pt(2)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.2)


def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    create_styles(doc)

    # =========================================================================
    # TITLE
    # =========================================================================
    doc.add_paragraph()  # Spacer

    title = doc.add_paragraph('PR Review Checklist Template', style='Title')

    subtitle = doc.add_paragraph(style='Subtitle')
    subtitle.clear()
    run = subtitle.add_run('for Claude-Generated Code')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.bold = True

    intro = doc.add_paragraph(
        'Use this checklist when reviewing pull requests that contain code generated '
        'or modified by Claude Code. Copy the relevant sections into your PR description '
        'or use as a reviewer guide.'
    )
    intro.paragraph_format.space_after = Pt(12)

    add_horizontal_rule(doc)

    # =========================================================================
    # QUICK CHECKLIST
    # =========================================================================
    doc.add_heading('Quick Checklist (Copy into PR Description)', level=1)

    # -- Correctness --
    add_section_label(doc, 'Correctness')
    add_checkbox_item(doc, 'All imports exist in installed package versions')
    add_checkbox_item(doc, 'API endpoints referenced in frontend exist in backend urls.py')
    add_checkbox_item(doc, 'Model fields and properties match the actual schema')
    add_checkbox_item(doc, 'No invented methods, mixins, or classes from external packages')
    add_checkbox_item(doc, "Framework APIs match the project's framework version")

    # -- Code Quality --
    add_section_label(doc, 'Code Quality')
    add_checkbox_item(doc, 'Follows existing project patterns and conventions')
    add_checkbox_item(doc, 'No unnecessary changes to files outside the scope of the PR')
    add_checkbox_item(doc, 'No hardcoded values that should be configuration or constants')
    add_checkbox_item(doc, 'Edge cases handled (null, empty, invalid input)')
    add_checkbox_item(doc, 'No duplicated logic that already exists in the codebase')

    # -- Testing --
    add_section_label(doc, 'Testing')
    add_checkbox_item(doc, 'Existing tests still pass')
    add_checkbox_item(doc, 'New tests cover the added/changed functionality')
    add_checkbox_item(doc, 'Tests verify at least one happy path and one failure case')
    add_checkbox_item(doc, 'Test assertions are meaningful (not just "doesn\'t throw")')

    # -- Security --
    add_section_label(doc, 'Security')
    add_checkbox_item(doc, 'No secrets, API keys, or credentials in committed code')
    add_checkbox_item(doc, 'Input validation present at system boundaries')
    add_checkbox_item(doc, 'No SQL injection, XSS, or command injection vectors')
    add_checkbox_item(doc, 'Authentication/authorization checks in place for new endpoints')

    # -- Integration --
    add_section_label(doc, 'Integration')
    add_checkbox_item(doc, 'Database migrations are correct and reversible (if applicable)')
    add_checkbox_item(doc, 'API contract changes are reflected in both backend and frontend')
    add_checkbox_item(doc, 'No breaking changes to existing interfaces without migration plan')
    add_checkbox_item(doc, 'Documentation updated if behavior changes')

    add_horizontal_rule(doc)

    # =========================================================================
    # EXTENDED CHECKLIST
    # =========================================================================
    doc.add_heading('Extended Checklist for Detailed Reviews', level=1)

    # -- Django / Backend --
    doc.add_heading('Django / Backend Specific', level=2)

    add_section_label(doc, 'Backend Review')
    add_checkbox_item(doc, 'Model changes have corresponding migrations')
    add_checkbox_item(doc, 'Serializer fields match model fields (names, types, read-only status)')
    add_checkbox_item(doc, 'ViewSet/View permissions are set correctly (AllowAdmin, AllowManager, etc.)')
    add_checkbox_item(doc, 'QuerySets are filtered appropriately (no unscoped .all() in user-facing views)')
    add_checkbox_item(doc, 'ForeignKey on_delete behavior is intentional')
    add_checkbox_item(doc, "New URL patterns don't conflict with existing ones")
    add_checkbox_item(doc, 'Admin registration updated if new models are added')

    # -- React / Frontend --
    doc.add_heading('React / Frontend Specific', level=2)

    add_section_label(doc, 'Frontend Review')
    add_checkbox_item(doc, "Components use hooks/patterns consistent with the project's React version")
    add_checkbox_item(doc, 'API calls use correct endpoints (verified against backend urls.py)')
    add_checkbox_item(doc, 'Error states are handled (loading, error, empty data)')
    add_checkbox_item(doc, 'State management follows project patterns (Redux, Context, etc.)')
    add_checkbox_item(doc, 'No direct DOM manipulation (use React patterns)')
    add_checkbox_item(doc, 'Translations/i18n keys added if user-facing strings are introduced')
    add_checkbox_item(doc, 'Component is used in a route or parent component (not orphaned)')

    add_horizontal_rule(doc)

    # =========================================================================
    # HOW TO USE THIS TEMPLATE
    # =========================================================================
    doc.add_heading('How to Use This Template', level=1)

    # -- For PR Authors --
    doc.add_heading('For PR Authors (using Claude Code)', level=2)

    add_numbered_item(doc, 1, 'After Claude creates the PR, ', bold_prefix='')
    p = doc.paragraphs[-1]
    run = p.add_run('add the Quick Checklist')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(' to the PR description')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    add_numbered_item(doc, 2, 'Go through each item yourself before requesting review')
    add_numbered_item(doc, 3, 'Check the items you\'ve verified; leave unchecked items for the reviewer')
    add_numbered_item(doc, 4, 'Add a note for any items that don\'t apply:')

    # Example of N/A notation
    p = doc.add_paragraph()
    p.style = doc.styles['Custom Bullet']
    p.paragraph_format.left_indent = Inches(0.75)
    p.clear()
    run = p.add_run('\u2611  N/A \u2014 no new endpoints')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -- For PR Reviewers --
    doc.add_heading('For PR Reviewers', level=2)

    add_numbered_item(doc, 1, '', bold_prefix='')
    p = doc.paragraphs[-1]
    run = p.add_run('Start with the ')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run('Correctness')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(" section \u2014 these catch Claude's most common errors")
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    add_numbered_item(doc, 2, 'Pay special attention to:')
    add_sub_bullet(doc, " \u2014 Claude frequently invents packages or uses wrong versions", bold_prefix='Imports')
    add_sub_bullet(doc, " \u2014 Claude may reference endpoints that don't exist", bold_prefix='API endpoints')
    add_sub_bullet(doc, " \u2014 Claude defaults to patterns from newer versions", bold_prefix='Framework version')

    add_numbered_item(doc, 3, 'Use the Extended Checklist for larger PRs or PRs touching critical paths')
    add_numbered_item(doc, 4, "If you find a hallucination, add it to the team's CLAUDE.md as a constraint")

    # -- For the Team --
    doc.add_heading('For the Team (continuous improvement)', level=2)

    doc.add_paragraph('After each sprint or review cycle:')
    add_bullet(doc, 'Collect common Claude mistakes from PR reviews')
    add_bullet(doc, 'Add preventive constraints to CLAUDE.md')
    add_bullet(doc, 'Update .claude/settings.json deny rules if Claude keeps trying risky operations')
    add_bullet(doc, 'Share useful prompts that produced good results')

    add_horizontal_rule(doc)

    # =========================================================================
    # COMMON CLAUDE MISTAKES TABLE
    # =========================================================================
    doc.add_heading('Common Claude Mistakes to Watch For', level=1)

    mistakes_table = add_styled_table(
        doc,
        ['Category', 'What to Look For', 'Example'],
        [
            [
                'Phantom imports',
                'Imports from packages not in your dependencies',
                'from rest_framework.mixins import BulkCreateMixin'
            ],
            [
                'Version mismatch',
                'Uses APIs from newer framework versions',
                'useActionState (React 19) in a React 16 project'
            ],
            [
                'Invented endpoints',
                "Frontend calls API endpoints that don't exist",
                "axios.get('/api/users/permissions/')"
            ],
            [
                'Wrong relationships',
                'Misunderstands model relationships',
                'Direct Client \u2192 Invoice instead of Client \u2192 Contract \u2192 Invoice'
            ],
            [
                'Over-engineering',
                'Adds abstraction layers not asked for',
                'Creating a utility class for a one-time operation'
            ],
            [
                'Missing edge cases',
                'Happy path only, no null/error handling',
                'Property crashes on None foreign key'
            ],
            [
                'Scope creep',
                'Modifies files outside the task scope',
                '"While I was there, I also refactored..."'
            ],
        ],
        col_widths=[1.5, 2.3, 2.7]
    )

    # Make the example column use monospace
    for row in mistakes_table.rows[1:]:
        for run in row.cells[2].paragraphs[0].runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x88, 0x44, 0x44)

    add_horizontal_rule(doc)

    # =========================================================================
    # PR DESCRIPTION TEMPLATE
    # =========================================================================
    doc.add_heading('PR Description Template', level=1)

    doc.add_paragraph(
        'Use this as a complete PR description template for Claude-assisted work:'
    )

    # Render the template as a styled block rather than raw code,
    # since it's meant to be copied and filled in.

    # Summary section
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run('Summary')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('[1\u20133 sentences: what was changed and why]')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Calibri'

    # Changes section
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run('Changes')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.name = 'Calibri'

    for placeholder in ['[File 1]: [what changed]', '[File 2]: [what changed]', '[File 3]: [what changed]']:
        p = doc.add_paragraph()
        p.style = doc.styles['Custom Bullet']
        p.clear()
        run = p.add_run('\u2022  ')
        run.font.size = Pt(10)
        run = p.add_run(placeholder)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = 'Calibri'

    # How it was built section
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run('How it was built')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.name = 'Calibri'

    add_checked_item(doc, 'Planned in Claude Code Plan Mode')
    add_checked_item(doc, 'Implemented with Claude Code')
    add_checked_item(doc, 'Reviewed diffs manually before committing')
    add_checked_item(doc, 'Tests pass')

    # Testing section
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run('Testing')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.name = 'Calibri'

    for placeholder in ['[Describe how this was tested]', '[Include test commands run]']:
        p = doc.add_paragraph()
        p.style = doc.styles['Custom Bullet']
        p.clear()
        run = p.add_run('\u2022  ')
        run.font.size = Pt(10)
        run = p.add_run(placeholder)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = 'Calibri'

    # Review Checklist section within template
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run('Review Checklist')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.name = 'Calibri'

    add_checkbox_item(doc, 'Imports are real and version-compatible')
    add_checkbox_item(doc, 'API endpoints exist in the backend')
    add_checkbox_item(doc, 'Model changes have migrations (if applicable)')
    add_checkbox_item(doc, 'Edge cases handled')
    add_checkbox_item(doc, 'No scope creep \u2014 only requested changes')
    add_checkbox_item(doc, 'Existing tests pass')
    add_checkbox_item(doc, 'New functionality has test coverage')

    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(20)
    run = footer_p.add_run('\u2014 End of Template \u2014')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    return doc


if __name__ == '__main__':
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, '07_PR_Review_Checklist_Template.docx')

    doc = build_document()
    doc.save(output_path)
    print(f'Generated: {output_path}')
