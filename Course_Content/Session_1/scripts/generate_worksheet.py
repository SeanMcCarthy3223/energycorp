#!/usr/bin/env python3
"""Generate a professional Word document from the Student Worksheet markdown."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "2B579A")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    # Style data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

def add_code_block(doc, code_text):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.style = doc.styles['Code Block']
    p.clear()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    return p

def add_callout(doc, text, style='Callout'):
    """Add a callout/tip box."""
    p = doc.add_paragraph()
    p.style = doc.styles[style]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    return p

def add_checkbox_item(doc, text, indent_level=0):
    """Add a checkbox list item."""
    p = doc.add_paragraph()
    p.style = doc.styles['Checkbox']
    p.clear()
    # Add checkbox
    run = p.add_run('\u2610  ')  # Unicode empty checkbox
    run.font.size = Pt(11)
    # Add text
    run = p.add_run(text)
    run.font.size = Pt(10)
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * indent_level)
    return p

def add_answer_line(doc, label="Your answer:"):
    """Add an answer line for student responses."""
    p = doc.add_paragraph()
    p.style = doc.styles['Answer Line']
    run = p.add_run(f"{label} ")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run = p.add_run("_" * 70)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    return p

def create_styles(doc):
    """Create custom styles for the document."""
    styles = doc.styles

    # --- Title style ---
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(4)

    # --- Subtitle style ---
    subtitle = styles['Subtitle']
    subtitle.font.name = 'Calibri'
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)

    # --- Heading 1 (Exercise headers) ---
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    # Add bottom border
    pPr = h1.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="2B579A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # --- Heading 2 (Sub-exercise headers) ---
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # --- Heading 3 ---
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # --- Normal body text ---
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # --- Code Block style ---
    code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.left_indent = Inches(0.3)
    code_style.paragraph_format.right_indent = Inches(0.3)
    # Dark background via shading
    pPr = code_style.element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E" w:val="clear"/>')
    pPr.append(shading)

    # --- Callout style ---
    callout_style = styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH)
    callout_style.font.name = 'Calibri'
    callout_style.font.size = Pt(10)
    callout_style.font.italic = True
    callout_style.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    callout_style.paragraph_format.left_indent = Inches(0.4)
    callout_style.paragraph_format.space_before = Pt(6)
    callout_style.paragraph_format.space_after = Pt(6)
    # Light blue-gray left border and background
    pPr = callout_style.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="2B579A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF3FA" w:val="clear"/>')
    pPr.append(shading)

    # --- Checkbox style ---
    checkbox_style = styles.add_style('Checkbox', WD_STYLE_TYPE.PARAGRAPH)
    checkbox_style.font.name = 'Calibri'
    checkbox_style.font.size = Pt(10)
    checkbox_style.paragraph_format.space_before = Pt(2)
    checkbox_style.paragraph_format.space_after = Pt(2)
    checkbox_style.paragraph_format.left_indent = Inches(0.3)

    # --- Answer Line style ---
    answer_style = styles.add_style('Answer Line', WD_STYLE_TYPE.PARAGRAPH)
    answer_style.font.name = 'Calibri'
    answer_style.font.size = Pt(10)
    answer_style.paragraph_format.space_before = Pt(4)
    answer_style.paragraph_format.space_after = Pt(10)
    answer_style.paragraph_format.left_indent = Inches(0.4)

    # --- Bullet style ---
    bullet = styles.add_style('Custom Bullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet.font.name = 'Calibri'
    bullet.font.size = Pt(10)
    bullet.paragraph_format.space_before = Pt(2)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.2)

    # --- Tips style ---
    tips_style = styles.add_style('Tips', WD_STYLE_TYPE.PARAGRAPH)
    tips_style.font.name = 'Calibri'
    tips_style.font.size = Pt(10)
    tips_style.paragraph_format.left_indent = Inches(0.3)
    tips_style.paragraph_format.space_before = Pt(2)
    tips_style.paragraph_format.space_after = Pt(2)


def add_horizontal_rule(doc):
    """Add a horizontal rule / divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_prefix=None):
    """Add a styled bullet point."""
    p = doc.add_paragraph()
    p.style = doc.styles['Custom Bullet']
    p.clear()
    run = p.add_run('\u2022  ')
    run.font.size = Pt(10)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_tip_bullet(doc, text, bold_prefix=None):
    """Add a tip bullet point."""
    p = doc.add_paragraph()
    p.style = doc.styles['Tips']
    p.clear()
    run = p.add_run('\u2022  ')
    run.font.size = Pt(10)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def build_document():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    create_styles(doc)

    # =========================================================================
    # TITLE PAGE AREA
    # =========================================================================
    doc.add_paragraph()  # Spacer

    title = doc.add_paragraph('Hands-on Lab', style='Title')

    subtitle = doc.add_paragraph(style='Subtitle')
    subtitle.clear()
    run = subtitle.add_run('Getting Started with Claude Code')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.bold = True

    # Duration and Goal
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Duration: ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = meta.add_run('25\u201335 minutes (independent work)')
    run.font.size = Pt(10.5)

    goal = doc.add_paragraph()
    goal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = goal.add_run('Goal: ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = goal.add_run(
        'Set up Claude Code for your project, create a CLAUDE.md, '
        'implement a small change, and submit a PR with a review checklist.'
    )
    run.font.size = Pt(10.5)
    goal.paragraph_format.space_after = Pt(16)

    add_horizontal_rule(doc)

    # =========================================================================
    # EXERCISE 1: Setup & Verification
    # =========================================================================
    doc.add_heading('Exercise 1: Setup & Verification (3 min)', level=1)

    doc.add_heading('1.1  Verify Installation', level=2)
    doc.add_paragraph('Open your terminal and run:')
    add_code_block(doc, 'claude --version')
    doc.add_paragraph(
        'If you see a version number, skip to 1.2. If not, install:'
    )
    add_code_block(doc, 'npm install -g @anthropic-ai/claude-code')
    doc.add_paragraph(
        'Then run claude and follow the authentication prompts.'
    )

    doc.add_heading('1.2  Navigate to Your Project', level=2)
    add_code_block(doc, 'cd /path/to/your/project')

    doc.add_heading('1.3  Launch Claude Code', level=2)
    add_code_block(doc, 'claude')
    doc.add_paragraph('Verify Claude detects your project. Ask it:')
    add_code_block(doc, 'What can you tell me about this project from looking at the file structure?')

    doc.add_paragraph(
        'Write down what Claude gets right and what it gets wrong:'
    )

    # Table: Correct vs Incorrect
    table = add_styled_table(
        doc,
        ['Correct', 'Incorrect or Missing'],
        [['', ''], ['', ''], ['', '']],
        col_widths=[3.0, 3.0]
    )

    doc.add_paragraph()  # small spacer
    add_callout(
        doc,
        'The incorrect/missing items tell you what your CLAUDE.md needs to cover.'
    )

    add_horizontal_rule(doc)

    # =========================================================================
    # EXERCISE 2: Create Your CLAUDE.md
    # =========================================================================
    doc.add_heading('Exercise 2: Create Your CLAUDE.md (8\u201310 min)', level=1)

    doc.add_heading('2.1  Start with a Template', level=2)
    doc.add_paragraph(
        'Create a file called CLAUDE.md in your project root. '
        'Use this structure as a starting point:'
    )

    template_code = """# CLAUDE.md

## Project Overview

[One to two sentences: what is this project, what does it do, who uses it]

## Tech Stack

- Backend: [framework, version, language]
- Frontend: [framework, version, language]
- Database: [database, local vs production differences]
- Key libraries: [list 3-5 most important dependencies]

## Build & Run Commands

### Backend

```bash
[exact command to install dependencies]
[exact command to run dev server]
[exact command to run tests]
[exact command to run linter, if applicable]
[exact command to create/apply migrations]
```

### Frontend

```bash
[exact command to install dependencies]
[exact command to run dev server]
[exact command to run tests]
[exact command to build for production]
```

## Architecture

### Backend

[Describe your app structure: what are the main modules/apps,
 what does each one do, what are the API prefixes]

### Frontend

[Describe your component structure: routing, state management,
 key patterns]

## Key Model Relationships

[Describe the main data model chain(s).
 Example: User \u2192 Order \u2192 OrderItem \u2192 Product]

## Important Constraints

- [Anything Claude should NOT do or assume]
- [Version-specific limitations]
- [Team conventions that aren't obvious from the code]"""

    add_code_block(doc, template_code)

    doc.add_heading('2.2  Fill In the Template', level=2)
    doc.add_paragraph(
        'Use your knowledge of the project to fill in each section. '
        'Aim for 50\u201380 lines total.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Tips:')
    run.bold = True
    run.font.size = Pt(10.5)

    add_tip_bullet(doc, 'Django 4.2', bold_prefix='Be specific: ')
    p = doc.add_paragraph()
    p.style = doc.styles['Tips']
    p.clear()
    run = p.add_run('    not just ')
    run.font.size = Pt(10)
    run = p.add_run('Django')
    run.font.italic = True
    run.font.size = Pt(10)

    add_tip_bullet(doc, 'python manage.py test apps/', bold_prefix='Use exact commands: ')
    p = doc.add_paragraph()
    p.style = doc.styles['Tips']
    p.clear()
    run = p.add_run('    not ')
    run.font.size = Pt(10)
    run = p.add_run('run the tests')
    run.font.italic = True
    run.font.size = Pt(10)

    add_tip_bullet(doc, 'Include your actual app/module names and their responsibilities')
    add_tip_bullet(doc, 'List the 2\u20133 most important model relationships')

    doc.add_heading('2.3  Verify with Claude', level=2)
    doc.add_paragraph(
        'Save your CLAUDE.md, then restart Claude Code (exit and re-launch claude). Ask:'
    )
    add_code_block(
        doc,
        'What does the CLAUDE.md tell you about this project? Is there anything\n'
        "you'd recommend adding to make it more useful?"
    )
    doc.add_paragraph(
        'Claude may suggest additions. Add any that make sense, but remember: '
        'lean is better than comprehensive. Only add what will prevent real mistakes.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Self-check \u2014 Your CLAUDE.md should answer these questions:')
    run.bold = True
    run.font.size = Pt(10.5)

    add_checkbox_item(doc, 'What does this project do? (1\u20132 sentences)')
    add_checkbox_item(doc, "What's the tech stack? (with versions)")
    add_checkbox_item(doc, 'How do I run the tests? (exact command)')
    add_checkbox_item(doc, 'How is the code organized? (app/module structure)')
    add_checkbox_item(doc, 'How do the main models relate to each other?')

    add_horizontal_rule(doc)

    # =========================================================================
    # EXERCISE 3: Configure settings.json
    # =========================================================================
    doc.add_heading('Exercise 3: Configure settings.json (3 min)', level=1)

    doc.add_heading('3.1  Create the Settings Directory', level=2)
    add_code_block(doc, 'mkdir -p .claude')

    doc.add_heading('3.2  Create Project Settings', level=2)
    doc.add_paragraph(
        'Create .claude/settings.json with appropriate permissions for your project. '
        'Adapt this template:'
    )

    settings_code = """{
  "permissions": {
    "allow": [
      "Bash([your test command here])",
      "Bash([your lint command here])",
      "Bash([your build command here])"
    ],
    "deny": [
      "Bash(rm -rf *)",
      "Bash(*DROP TABLE*)",
      "Bash(*--force*)",
      "Bash(*manage.py flush*)"
    ]
  }
}"""
    add_code_block(doc, settings_code)

    p = doc.add_paragraph()
    run = p.add_run('Replace the bracketed items ')
    run.font.size = Pt(10.5)
    run = p.add_run('with your actual commands. For example:')
    run.font.size = Pt(10.5)

    add_bullet(doc, ' \u2014 allow running Django tests', bold_prefix='"Bash(python manage.py test)"')
    add_bullet(doc, ' \u2014 allow running frontend tests', bold_prefix='"Bash(npm test)"')
    add_bullet(doc, ' \u2014 allow running the linter', bold_prefix='"Bash(npx eslint .)"')

    doc.add_heading('3.3  (Optional) Create Local Settings', level=2)
    doc.add_paragraph(
        'If you want personal overrides, create .claude/settings.local.json:'
    )
    add_code_block(doc, '{\n  "model": "sonnet"\n}')
    doc.add_paragraph(
        'Add .claude/settings.local.json to your .gitignore if it\'s not already there.'
    )

    add_horizontal_rule(doc)

    # =========================================================================
    # EXERCISE 4: Implement a Small Change
    # =========================================================================
    doc.add_heading('Exercise 4: Implement a Small Change (10\u201312 min)', level=1)

    doc.add_heading('4.1  Choose Your Change', level=2)
    doc.add_paragraph(
        'Pick ONE of the following (or propose your own \u2014 something small and self-contained):'
    )

    p = doc.add_paragraph()
    run = p.add_run('Backend options:')
    run.bold = True
    run.font.size = Pt(10.5)

    add_bullet(doc, 'Add a new read-only field or property to an existing model')
    add_bullet(doc, 'Add a new serializer field that computes a value from existing data')
    add_bullet(doc, 'Add a new filter or search parameter to an existing API endpoint')
    add_bullet(doc, 'Write a utility function for a common operation in your codebase')

    p = doc.add_paragraph()
    run = p.add_run('Frontend options:')
    run.bold = True
    run.font.size = Pt(10.5)

    add_bullet(doc, 'Add a display component that formats existing data in a new way')
    add_bullet(doc, 'Add a new column to an existing table/list view')
    add_bullet(doc, 'Create a small reusable component (e.g., a status badge, formatted date display)')
    add_bullet(doc, 'Add a loading state to a component that doesn\'t have one')

    doc.add_heading('4.2  Plan First', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Enter Plan Mode: ')
    run.font.size = Pt(10.5)
    run = p.add_run('press Shift+Tab twice (or type /plan).')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    doc.add_paragraph('Ask Claude to plan your change:')
    add_code_block(
        doc,
        'Plan how to [describe your change]. Tell me which files need to change\n'
        'and in what order.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Review the plan. Verify:')
    run.bold = True
    run.font.size = Pt(10.5)

    add_checkbox_item(doc, 'Does it reference real files in your project?')
    add_checkbox_item(doc, 'Does the order of operations make sense?')
    add_checkbox_item(doc, 'Does it mention running tests?')
    add_checkbox_item(doc, 'Does it avoid unnecessary changes?')

    doc.add_heading('4.3  Implement', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Switch to Normal Mode: ')
    run.font.size = Pt(10.5)
    run = p.add_run('press Shift+Tab once.')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    doc.add_paragraph('Ask Claude to implement:')
    add_code_block(doc, 'Implement the plan: [brief description of the change]')

    doc.add_heading('4.4  Review the Code', level=2)
    doc.add_paragraph('Before accepting, check:')
    add_checkbox_item(doc, 'Are all imports real? (No phantom packages or methods)')
    add_checkbox_item(doc, "Does the code match your project's patterns? (Naming conventions, file organization)")
    add_checkbox_item(doc, 'Are there any hardcoded values that should be configurable?')
    add_checkbox_item(doc, 'Does it handle edge cases? (null values, empty strings, etc.)')

    add_horizontal_rule(doc)

    # =========================================================================
    # EXERCISE 5: Validate
    # =========================================================================
    doc.add_heading('Exercise 5: Validate (3 min)', level=1)

    doc.add_heading('5.1  Run Tests', level=2)
    doc.add_paragraph('Ask Claude:')
    add_code_block(doc, 'Run the tests to verify nothing is broken.')
    doc.add_paragraph('Or run them yourself:')
    add_code_block(doc, '[your test command]')

    p = doc.add_paragraph()
    run = p.add_run('Test results:')
    run.bold = True
    run.font.size = Pt(10.5)

    add_checkbox_item(doc, 'All tests pass')
    add_checkbox_item(doc, 'If tests fail, I gave Claude the error and it fixed the issue')
    add_checkbox_item(doc, 'If it took more than 2 corrections, I used /clear and started fresh')

    doc.add_heading('5.2  Manual Review', level=2)
    doc.add_paragraph('Check the diff manually:')
    add_code_block(doc, 'git diff')

    doc.add_paragraph('Verify:')
    add_checkbox_item(doc, 'Only the intended files were changed')
    add_checkbox_item(doc, 'No unrelated changes snuck in')
    add_checkbox_item(doc, 'The code does what was requested')

    add_horizontal_rule(doc)

    # =========================================================================
    # EXERCISE 6: Create a PR with Review Checklist
    # =========================================================================
    doc.add_heading('Exercise 6: Create a PR with Review Checklist (5 min)', level=1)

    doc.add_heading('6.1  Create the PR', level=2)
    doc.add_paragraph('Ask Claude:')
    add_code_block(
        doc,
        'Create a new branch called "feature/[short-description]", commit the\n'
        'changes, and create a PR. Include a summary and a review checklist\n'
        'for the reviewer.'
    )

    doc.add_paragraph(
        "If Claude can't create a PR directly (e.g., no GitHub CLI), do it manually:"
    )
    add_code_block(
        doc,
        'git checkout -b feature/[short-description]\n'
        'git add [changed files]\n'
        'git commit -m "Add [description of change]"\n'
        'git push -u origin feature/[short-description]'
    )
    doc.add_paragraph(
        "Then create the PR through your platform's web interface."
    )

    doc.add_heading('6.2  Add the Review Checklist', level=2)
    doc.add_paragraph(
        "If Claude didn't generate one, add this checklist to your PR description (adapt as needed):"
    )

    # Review checklist as a structured section
    p = doc.add_paragraph()
    run = p.add_run('Code Quality')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    add_checkbox_item(doc, 'All imports reference real, installed packages')
    add_checkbox_item(doc, 'Code follows project naming conventions')
    add_checkbox_item(doc, 'No hardcoded values that should be constants or config')
    add_checkbox_item(doc, 'Edge cases handled (null, empty, unexpected input)')

    p = doc.add_paragraph()
    run = p.add_run('Testing')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    add_checkbox_item(doc, 'Existing tests still pass')
    add_checkbox_item(doc, 'New functionality has test coverage (or justification for skipping)')
    add_checkbox_item(doc, 'Tests cover happy path and at least one error case')

    p = doc.add_paragraph()
    run = p.add_run('Integration')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    add_checkbox_item(doc, 'Changes are backward compatible')
    add_checkbox_item(doc, 'API changes are reflected in serializers')
    add_checkbox_item(doc, 'Frontend changes match backend data contracts')
    add_checkbox_item(doc, 'No unnecessary files modified')

    p = doc.add_paragraph()
    run = p.add_run('Claude-Specific Checks')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    add_checkbox_item(doc, 'No hallucinated imports or API calls')
    add_checkbox_item(doc, "Code matches the project's framework version")
    add_checkbox_item(doc, "Generated code doesn't duplicate existing utilities")
    add_checkbox_item(doc, "Review the full diff, not just Claude's summary")

    add_horizontal_rule(doc)

    # =========================================================================
    # SELF-ASSESSMENT
    # =========================================================================
    doc.add_heading('Self-Assessment', level=1)
    doc.add_paragraph(
        "Answer these questions honestly. They're for your own knowledge and retention."
    )

    doc.add_heading('Understanding', level=2)

    p = doc.add_paragraph()
    run = p.add_run('1. ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run('What is the purpose of CLAUDE.md?')
    run.font.size = Pt(10.5)
    add_answer_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('2. ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run(
        'What are the five tiers of settings.json, from highest to lowest priority?'
    )
    run.font.size = Pt(10.5)
    add_answer_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('3. ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run(
        'Why should you use Plan Mode before implementing a non-trivial change?'
    )
    run.font.size = Pt(10.5)
    add_answer_line(doc)

    doc.add_heading('Experience', level=2)

    p = doc.add_paragraph()
    run = p.add_run('4. ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run('Did Claude hallucinate anything during your session? If so, what?')
    run.font.size = Pt(10.5)
    add_answer_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('5. ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run('Did you hit the Correction Spiral? What did you do?')
    run.font.size = Pt(10.5)
    add_answer_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('6. ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run(
        "What's one thing you'd add to your CLAUDE.md based on today's experience?"
    )
    run.font.size = Pt(10.5)
    add_answer_line(doc)

    doc.add_heading('Confidence', level=2)
    doc.add_paragraph('Rate your comfort level (1 = not comfortable, 5 = very comfortable):')

    confidence_table = add_styled_table(
        doc,
        ['Skill', 'Rating (1\u20135)'],
        [
            ['Launching and using Claude Code CLI', ''],
            ['Writing a useful CLAUDE.md', ''],
            ['Using Plan Mode before implementing', ''],
            ['Recognizing hallucinations', ''],
            ["Creating a PR with Claude's help", ''],
            ['Knowing when to /clear and start fresh', ''],
        ],
        col_widths=[4.5, 1.5]
    )
    # Center the rating column
    for row in confidence_table.rows[1:]:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_horizontal_rule(doc)

    # =========================================================================
    # WHAT TO BRING TO SESSION 2
    # =========================================================================
    doc.add_heading('What to Bring to Session 2', level=1)
    add_bullet(doc, "Your CLAUDE.md (we'll review and improve them)")
    add_bullet(doc, 'Notes on any Claude mistakes or frustrations you encountered')
    add_bullet(doc, 'One question about Claude Code you want answered')

    add_horizontal_rule(doc)

    # =========================================================================
    # QUICK REFERENCE
    # =========================================================================
    doc.add_heading('Quick Reference', level=1)

    ref_table = add_styled_table(
        doc,
        ['Action', 'How'],
        [
            ['Launch Claude Code', 'claude in your project directory'],
            ['Switch to Plan Mode', 'Shift+Tab twice, or /plan'],
            ['Switch to Normal Mode', 'Shift+Tab'],
            ['Compress context', '/compact'],
            ['Clear session', '/clear'],
            ['Reference a file', '@path/to/file'],
            ['Switch model', '/model'],
            ['Get help', '/help'],
        ],
        col_widths=[2.5, 3.5]
    )

    # Make the "How" column use monospace font
    for row in ref_table.rows[1:]:
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph()  # Spacer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(20)
    run = footer_p.add_run('\u2014 End of Worksheet \u2014')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    return doc


if __name__ == '__main__':
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, '06_Student_Worksheet.docx')

    doc = build_document()
    doc.save(output_path)
    print(f'Generated: {output_path}')
